"""
VEM Marketing Record generator.

Recreates the structure of VEM_Marketing_Record_Template_0008.docx (the
firm's SEC Marketing Rule recordkeeping template) and fills it in from a
plain Python dict, so a new record can be produced automatically every time
a marketing communication (blog post, social post, etc.) is published.

Usage:
    python scripts/generate_marketing_record.py

Edit the RECORDS list at the bottom of this file with one dict per record,
then run. Each record is written to <out_dir>/<Marketing Record ID>.docx.

Any field left as "TBD" in a filled record is intentional — it means this
script did not have a confirmed answer and a human needs to fill it in
before the record is considered complete. Never guess a value for a
compliance field (Compliance Reviewer, Approval Date, Published By, etc.)
just to avoid a blank.
"""
import os
import docx
from docx.shared import Pt, Inches

TBD = "TBD"


def add_heading(doc, text, level=2):
    doc.add_heading(text, level=level)


def add_field_table(doc, rows):
    """rows: list of (field, value) tuples. First row is always the header."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (field, value) in enumerate(rows):
        table.cell(i, 0).text = str(field)
        table.cell(i, 1).text = str(value)
    return table


def add_list_table(doc, header, items):
    """Single-column table: header row + one row per item (blank rows padded to at least 5, matching the official template's 6-row Source Document table)."""
    rows = [header] + list(items)
    while len(rows) < 6:
        rows.append("")
    table = doc.add_table(rows=len(rows), cols=1)
    table.style = "Table Grid"
    for i, val in enumerate(rows):
        table.cell(i, 0).text = str(val)
    return table


def add_log_table(doc, header, entries):
    """entries: list of (col1, col2, col3) tuples."""
    rows = [header] + list(entries)
    while len(rows) < 5:
        rows.append(("", "", ""))
    table = doc.add_table(rows=len(rows), cols=3)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            table.cell(i, j).text = str(val)
    return table


def build_record(data):
    """data: dict, see RECORDS below for the expected shape."""
    doc = docx.Document()

    doc.add_heading("VEM Marketing Record", level=1)
    doc.add_paragraph("Complete one Marketing Record for every published marketing communication.")

    add_heading(doc, "1. Record Information")
    add_field_table(doc, [
        ("Field", "Value"),
        ("Marketing Record ID", data.get("record_id", "")),
        ("Post Title", data.get("post_title", "")),
        ("Status (Draft/Approved/Published/Archived)", data.get("status", "")),
        ("Date Created", data.get("date_created", "")),
        ("Publish Date", data.get("publish_date", "")),
        ("Archive Date", data.get("archive_date", "")),
    ])

    add_heading(doc, "2. Post Metadata")
    add_field_table(doc, [
        ("Field", "Value"),
        ("Platform(s)", data.get("platforms", "")),
        ("Content Category", data.get("content_category", "")),
        ("Post Type", data.get("post_type", "")),
        ("Author", data.get("author", "")),
        ("Compliance Reviewer", data.get("compliance_reviewer", "")),
        ("Approval Date", data.get("approval_date", "")),
        ("Published By", data.get("published_by", "")),
        ("Published URL", data.get("published_url", "")),
        ("Screenshot Saved (Y/N)", data.get("screenshot_saved", "")),
        ("Archive Folder", data.get("archive_folder", "")),
    ])

    add_heading(doc, "3. Media Used")
    doc.add_paragraph("Insert a thumbnail/screenshot of the final approved creative below.")
    screenshot_path = data.get("screenshot_path")
    if screenshot_path and os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(4))
    else:
        doc.add_paragraph("[ Paste Image Here ]")
    media_filenames = data.get("media_filenames")
    if media_filenames:
        doc.add_paragraph("Media Filename: " + media_filenames)
    else:
        doc.add_paragraph("Media Filename: ________________________________")

    add_heading(doc, "4. Final Approved Caption (Published Version)")
    doc.add_paragraph("Paste the exact caption that was approved and published.")
    for para in data.get("caption_paragraphs", []):
        doc.add_paragraph(para)

    add_heading(doc, "5. Disclosure Used")
    doc.add_paragraph("Paste the exact disclosure that appeared with the published post.")
    disclosure = data.get("disclosure", "")
    if disclosure:
        doc.add_paragraph(disclosure)

    add_heading(doc, "6. Supporting Documentation")
    add_list_table(doc, "Source Document", data.get("sources", []))

    add_heading(doc, "7. Approval Log")
    add_log_table(doc, ("Step", "Person", "Date"), data.get("approval_log", [
        ("Draft Completed", "", ""),
        ("Compliance Review", "", ""),
        ("Approved", "", ""),
        ("Published", "", ""),
    ]))

    add_heading(doc, "8. Revision History")
    add_log_table(doc, ("Version", "Date", "Description of Changes"), data.get("revision_history", [
        ("", "", ""),
    ]))

    add_heading(doc, "9. Archive Checklist")
    for item in [
        "Final media file saved",
        "Final caption saved",
        "Live post screenshot saved",
        "Supporting documentation saved",
        "Approval documentation saved",
        "Marketing Archive Log updated",
        "Folder complete and archived",
    ]:
        doc.add_paragraph(u"☐ " + item)

    add_heading(doc, "10. Notes")
    doc.add_paragraph(data.get("notes", ""))

    # Fixed signature block — part of the official template, not per-record data.
    doc.add_paragraph("")
    doc.add_paragraph("John Vann CCO, CEO")

    return doc


if __name__ == "__main__":
    import sys
    print("Import this module and call build_record(data).save(path) — see fill_records.py for real usage.")
